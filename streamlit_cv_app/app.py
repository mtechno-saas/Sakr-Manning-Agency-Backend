import streamlit as st
import pandas as pd
import json
import time
import os
import re
import requests
import google.generativeai as genai
import PyPDF2
import docx
from io import BytesIO

# --- AI Extraction Logic ---
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(docx_file):
    """Extract text from a DOCX file using python-docx."""
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text

def parse_cv_with_ai(file_bytes, cv_text, api_key, target_rank="All Ranks", filename="", is_docx=False):
    genai.configure(api_key=api_key)
    model_name = None
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if 'flash' in m.name or 'pro' in m.name:
                    model_name = m.name
                    break
    except Exception as e:
        return {"error": f"API Error (Likely invalid or expired API Key): {str(e)}"}
                
    if not model_name:
        return {"error": "No supported text generation models found for this API key."}
        
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    You are an AI assistant that extracts ALL structured information from seafarer resumes/CVs.
    
    Target Rank to look for: {target_rank}
    Filename (use as a hint if the CV text is incomplete): {filename}
    
    Extract the following details from the {'document text below' if is_docx else 'attached PDF document'}.
    
    PERSONAL INFORMATION:
    - full_name: Full Name (if not found, infer from filename)
    - email: Email Address (look for Email, E-mail, Gmail, Mail, etc.)
    - phone: Phone Number (look for Phone, Mobile, Tel, Mob, Cell, WhatsApp, etc.)
    - rank: Most recent rank/position (if matches Target Rank '{target_rank}', output exactly that)
    - date_of_birth: Date of birth (format: YYYY-MM-DD)
    - nationality: Nationality
    - place_of_birth: Place of birth
    - marital_status: Marital Status (Single or Married)
    - height_cm: Height in cm (number only)
    - weight_kg: Weight in kg (number only)
    - address: Home address
    - country: Country
    - city: City
    
    PASSPORT:
    - passport_no: Passport number
    - passport_issue_date: Passport issue date (YYYY-MM-DD)
    - passport_expiry_date: Passport expiry date (YYYY-MM-DD)
    - passport_issued_by: Passport issuing authority
    - passport_place_of_issue: Passport place of issue
    
    SEAMAN BOOK:
    - seaman_book_no: Seaman book number
    - seaman_book_issue_date: Seaman book issue date (YYYY-MM-DD)
    - seaman_book_expiry_date: Seaman book expiry date (YYYY-MM-DD)
    - seaman_book_issued_by: Seaman book issuing authority
    - seaman_book_place_of_issue: Seaman book place of issue
    
    NEXT OF KIN / EMERGENCY CONTACT:
    - next_of_kin_full_name: Emergency contact name
    - next_of_kin_relationship: Relationship
    - next_of_kin_phone: Emergency contact phone
    - next_of_kin_address_country: Emergency contact address/country
    
    SEA SERVICE (list of previous ship experiences):
    - sea_services: Array of objects, each with: vessel_name, rank, signed_on (YYYY-MM-DD), signed_off (YYYY-MM-DD)
    
    CRITICAL RULES:
    - Use YYYY-MM-DD format for ALL dates
    - Use empty string "" for any field you cannot find (never use null)
    - Extract as many sea service records as you can find
    - Try your best to fill every field from the document
    
    Return ONLY a valid JSON object with this exact structure:
    {{
        "full_name": "",
        "email": "",
        "phone": "",
        "rank": "",
        "date_of_birth": "",
        "nationality": "",
        "place_of_birth": "",
        "marital_status": "",
        "height_cm": "",
        "weight_kg": "",
        "address": "",
        "country": "",
        "city": "",
        "passport_no": "",
        "passport_issue_date": "",
        "passport_expiry_date": "",
        "passport_issued_by": "",
        "passport_place_of_issue": "",
        "seaman_book_no": "",
        "seaman_book_issue_date": "",
        "seaman_book_expiry_date": "",
        "seaman_book_issued_by": "",
        "seaman_book_place_of_issue": "",
        "next_of_kin_full_name": "",
        "next_of_kin_relationship": "",
        "next_of_kin_phone": "",
        "next_of_kin_address_country": "",
        "sea_services": []
    }}
    """
    
    try:
        if is_docx:
            docx_prompt = prompt + f"\n\n--- DOCUMENT TEXT ---\n{cv_text}\n--- END OF DOCUMENT ---"
            response = model.generate_content(docx_prompt)
        else:
            response = model.generate_content([
                {"mime_type": "application/pdf", "data": file_bytes},
                prompt
            ])
        text_resp = response.text.strip()
        
        if text_resp.startswith("```json"):
            text_resp = text_resp[7:]
        elif text_resp.startswith("```"):
            text_resp = text_resp[3:]
            
        if text_resp.endswith("```"):
            text_resp = text_resp[:-3]
            
        return json.loads(text_resp.strip())
    except Exception as e:
        return {"error": str(e)}

# --- Django Integration Logic ---
def send_to_documents(extracted_data, file_obj, base_url, auth_token=""):
    """Step 1: Send extracted data and file to /api/documents/ endpoint."""
    file_obj.seek(0)
    
    file_ext = os.path.splitext(file_obj.name)[1].lower()
    if file_ext == '.docx':
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        mime_type = 'application/pdf'
    
    files = {
        'file': (file_obj.name, file_obj, mime_type)
    }
    
    data = {
        'name': extracted_data.get("full_name", ""),
        'email': extracted_data.get("email", ""),
        'phone': extracted_data.get("phone", ""),
        'position': extracted_data.get("rank", ""),
    }
    
    headers = {}
    if auth_token:
        headers['Authorization'] = f'Bearer {auth_token}'  
        
    try:
        url = base_url.rstrip('/') + '/api/documents/'
        response = requests.post(url, data=data, files=files, headers=headers)
        response.raise_for_status()
        
        try:
            resp_data = response.json()
        except ValueError:
            resp_data = {"status": "Success, but no JSON returned"}
            
        return True, resp_data
    except requests.exceptions.RequestException as e:
        err_msg = str(e)
        if hasattr(e, 'response') and e.response is not None:
            err_msg += f" | Body: {e.response.text}"
        return False, err_msg


def send_to_cv_submissions(extracted_data, file_obj, base_url, doc_response, auth_token=""):
    """Step 2: Send extracted data to /api/cv-submissions/ endpoint."""
    file_obj.seek(0)
    
    file_ext = os.path.splitext(file_obj.name)[1].lower()
    if file_ext == '.docx':
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        mime_type = 'application/pdf'
    
    user_id = doc_response.get('user')
    if not user_id:
        return False, "Could not get user ID from document response"
    
    files = {
        'cv_file': (file_obj.name, file_obj, mime_type)
    }
    
    data = {
        'user': user_id,
        'position': extracted_data.get("rank", ""),
        'user_email': extracted_data.get("email", ""),
        'status': 'Pending',
        'notes': 'Auto-submitted via AI CV Extractor',
    }
    
    headers = {}
    if auth_token:
        headers['Authorization'] = f'Bearer {auth_token}'  
        
    try:
        url = base_url.rstrip('/') + '/api/cv-submissions/'
        response = requests.post(url, data=data, files=files, headers=headers)
        response.raise_for_status()
        
        try:
            resp_data = response.json()
        except ValueError:
            resp_data = {"status": "Success, but no JSON returned"}
            
        return True, resp_data
    except requests.exceptions.RequestException as e:
        err_msg = str(e)
        if hasattr(e, 'response') and e.response is not None:
            err_msg += f" | Body: {e.response.text}"
        return False, err_msg


def update_user_profile(extracted_data, base_url, user_id, auth_token=""):
    """Step 3: Update the user profile with ALL extracted details via PATCH."""
    
    # Build the update payload — only include non-empty fields
    user_data = {}
    
    # Personal details
    field_map = {
        'date_of_birth': 'date_of_birth',
        'nationality': 'nationality',
        'place_of_birth': 'Place_Of_Birth',
        'marital_status': 'marital_status',
        'height_cm': 'Height_Cm',
        'weight_kg': 'Weight_Kg',
        'address': 'address',
        'country': 'country',
        'city': 'city',
        'phone': 'phone_number',
        # Passport
        'passport_no': 'passport_no',
        'passport_issue_date': 'passport_issue_date',
        'passport_expiry_date': 'passport_expiry_date',
        'passport_issued_by': 'passport_issued_by',
        'passport_place_of_issue': 'passport_place_of_issue',
        # Seaman Book
        'seaman_book_no': 'seaman_book_no',
        'seaman_book_issue_date': 'seaman_book_issue_date',
        'seaman_book_expiry_date': 'seaman_book_expiry_date',
        'seaman_book_issued_by': 'seaman_book_issued_by',
        'seaman_book_place_of_issue': 'seaman_book_place_of_issue',
        # Next of Kin
        'next_of_kin_full_name': 'next_of_kin_full_name',
        'next_of_kin_relationship': 'next_of_kin_relationship',
        'next_of_kin_phone': 'next_of_kin_phone',
        'next_of_kin_address_country': 'next_of_kin_address_country',
    }
    
    for ai_key, django_key in field_map.items():
        val = extracted_data.get(ai_key, "")
        if val and str(val).strip():
            # Convert height/weight to integers
            if django_key in ('Height_Cm', 'Weight_Kg'):
                try:
                    user_data[django_key] = int(float(str(val).strip()))
                except (ValueError, TypeError):
                    pass
            else:
                user_data[django_key] = str(val).strip()
    
    # Set the rank/position
    rank = extracted_data.get("rank", "")
    if rank:
        user_data['application_for_position'] = rank
        user_data['rank_ids'] = [rank]
    
    if not user_data:
        return True, {"status": "No additional data to update"}
    
    headers = {'Content-Type': 'application/json'}
    if auth_token:
        headers['Authorization'] = f'Bearer {auth_token}'
    
    try:
        url = base_url.rstrip('/') + f'/api/users/users/{user_id}/'
        response = requests.patch(url, json=user_data, headers=headers)
        response.raise_for_status()
        
        try:
            resp_data = response.json()
        except ValueError:
            resp_data = {"status": "Profile updated successfully"}
            
        return True, resp_data
    except requests.exceptions.RequestException as e:
        err_msg = str(e)
        if hasattr(e, 'response') and e.response is not None:
            err_msg += f" | Body: {e.response.text}"
        return False, err_msg


def add_sea_services(sea_services, base_url, user_id, auth_token=""):
    """Step 4: Add sea service records to the user profile."""
    if not sea_services:
        return True, {"status": "No sea service records to add"}
    
    headers = {'Content-Type': 'application/json'}
    if auth_token:
        headers['Authorization'] = f'Bearer {auth_token}'
    
    results = []
    for service in sea_services:
        data = {}
        if service.get('vessel_name'):
            data['vessel_name'] = service['vessel_name']
        if service.get('rank'):
            data['rank'] = service['rank']
        if service.get('signed_on'):
            data['signed_on'] = service['signed_on']
        if service.get('signed_off'):
            data['signed_off'] = service['signed_off']
        
        if not data.get('vessel_name'):
            continue
            
        try:
            url = base_url.rstrip('/') + f'/api/users/{user_id}/sea-services/'
            response = requests.post(url, json=data, headers=headers)
            response.raise_for_status()
            results.append({"vessel": data.get('vessel_name'), "status": "ok"})
        except requests.exceptions.RequestException as e:
            results.append({"vessel": data.get('vessel_name'), "status": str(e)})
    
    return True, results


# --- Streamlit UI ---
st.set_page_config(page_title="CV AI Extractor", page_icon="🤖", layout="wide")

st.title("🤖 AI-Powered CV Extractor (Full Profile)")
st.markdown("Upload CVs → Extract **all** details with AI → Save to **Documents** + **CV Submissions** + **User Profile**.")

# Sidebar for Configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Gemini API Key", type="password", help="Required for AI extraction.")
    st.divider()
    st.subheader("Django Connection")
    django_base_url = st.text_input("Django Backend URL", value="https://backend.sakrshipping.com", help="Base URL without trailing path")
    django_token = st.text_input("Django Auth Token (Optional)", type="password", help="If your API requires authentication.")

# Upload Area
st.subheader("📤 Upload Candidate CVs")

ranks = [
    "All Ranks", "Master / Captain", "Staff Captain", "Chief Officer / Chief Mate", "Second Officer",
    "Third Officer", "Dynamic Positioning Operator (DPO)", "ROV Supervisor",
    "Offshore Installation Manager", "Deck Cadet", "Bosun", "ABLE SEAFARER DECK",
    "Able Seaman (AB)", "Ordinary Seaman (OS)", "Carpenter", "Pumpman", "Crane Operator",
    "Water and Pool", "Security Guard", "Life Guard", "Upholsterer", "Doctor",
    "Hotel Director", "Assistant Hotel Director", "Purser", "Assistant Purser",
    "Food & Beverage Manager", "Executive Chef", "Chief Housekeeper", "Guest Services Manager",
    "Restaurant Manager", "Head Waiter", "Waiter", "F&B attendant", "Bartender",
    "Cabin Steward", "Laundryman", "Cook", "2nd Cook", "3rd Cook", "Assistant Cook",
    "Baker", "Assistant Baker", "Pastry", "Assistant pastry", "Butcher", "Steward",
    "Utility Galley", "Tour Expert", "Photographer", "Chief Engineer", "Second Engineer",
    "Third Engineer", "Fourth Engineer", "ETO", "2ND ETO", "3RD ETO", "ELECTRICAL ENGINEER",
    "Refrigeration Engineer", "HVAC Engineer", "Engine Cadet", "Gas Engineer",
    "Cargo Engineer", "Reliquefaction Engineer", "Able Seafarer Engine III/5", "Motorman",
    "Mechanic", "Oiler", "Wiper/Assistant Mechanic", "Fitter", "Welder",
    "Plumber", "Assistant Plumber", "Electrician",
    "2nd Electrician", "3rd Electrician", "Assistant Electrician", "Trainee Electrician",
    "AC Technician", "Senior Accommodation Repairman", "Junior Accommodation Repairman", "Other"
]
target_rank = st.selectbox("Filter by Target Rank", options=ranks)

uploaded_files = st.file_uploader("Drop PDF or DOCX files here", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    if not api_key:
        st.warning("⚠️ Please enter your Gemini API Key in the sidebar.")
    elif not django_base_url:
        st.warning("⚠️ Please enter the Django Backend URL in the sidebar.")
    else:
        if st.button("Extract & Send to Django", type="primary"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            successful_uploads = []
            
            for i, file in enumerate(uploaded_files):
                status_text.text(f"Processing {file.name} ({i+1}/{len(uploaded_files)})...")
                
                # Step A: Read raw bytes and extract text
                file.seek(0)
                file_bytes = file.read()
                file.seek(0)
                
                file_ext = os.path.splitext(file.name)[1].lower()
                is_docx = file_ext == '.docx'
                
                if is_docx:
                    cv_text = extract_text_from_docx(file)
                else:
                    cv_text = extract_text_from_pdf(file)
                
                # Step B: AI Parse — extract ALL details
                result = parse_cv_with_ai(file_bytes, cv_text, api_key, target_rank, filename=file.name, is_docx=is_docx)
                
                if "error" in result:
                    st.error(f"❌ Failed to parse {file.name}: {result['error']}")
                else:
                    # Fallback: name from filename
                    if not result.get("full_name"):
                        clean_name = re.sub(r'\.(pdf|docx)$', '', file.name, flags=re.IGNORECASE)
                        clean_name = re.sub(r'_Application|_CV|\d+', '', clean_name, flags=re.IGNORECASE)
                        clean_name = clean_name.replace('_', ' ').strip()
                        result["full_name"] = clean_name
                        st.info(f"ℹ️ Name fallback from filename: {clean_name}")

                    # Show AI extraction results
                    with st.expander(f"🔍 AI Extraction for {file.name}", expanded=False):
                        st.json(result)
                    
                    extracted_rank = result.get("rank", "Unknown")
                    if not extracted_rank or extracted_rank == "Unknown":
                        extracted_rank = target_rank if target_rank != "All Ranks" else "Unknown"
                        result["rank"] = extracted_rank
                    
                    if target_rank != "All Ranks" and extracted_rank.lower() != target_rank.lower():
                        st.warning(f"⏭️ Skipped {file.name}: Rank ({extracted_rank}) ≠ target ({target_rank})")
                    else:
                        # Step C: Send to /api/documents/
                        status_text.text(f"📄 Step 1/4: Saving document for {file.name}...")
                        doc_success, doc_response = send_to_documents(result, file, django_base_url, django_token)
                        
                        if doc_success:
                            st.success(f"✅ Step 1: Document saved")
                            user_id = doc_response.get('user')
                            
                            # Step D: Send to /api/cv-submissions/
                            status_text.text(f"📋 Step 2/4: Creating CV Submission for {file.name}...")
                            cv_success, cv_response = send_to_cv_submissions(result, file, django_base_url, doc_response, django_token)
                            
                            if cv_success:
                                st.success(f"✅ Step 2: CV Submission created")
                            else:
                                st.warning(f"⚠️ Step 2 failed: {cv_response}")
                            
                            # Step E: Update user profile with ALL extracted details
                            if user_id:
                                status_text.text(f"👤 Step 3/4: Updating user profile for {file.name}...")
                                profile_success, profile_response = update_user_profile(result, django_base_url, user_id, django_token)
                                
                                if profile_success:
                                    st.success(f"✅ Step 3: User profile updated")
                                else:
                                    st.warning(f"⚠️ Step 3 failed: {profile_response}")
                                
                                # Step F: Add sea service records
                                sea_services = result.get("sea_services", [])
                                if sea_services:
                                    status_text.text(f"🚢 Step 4/4: Adding {len(sea_services)} sea service records...")
                                    ss_success, ss_response = add_sea_services(sea_services, django_base_url, user_id, django_token)
                                    
                                    if ss_success:
                                        st.success(f"✅ Step 4: {len(sea_services)} sea service records added")
                                    else:
                                        st.warning(f"⚠️ Step 4 failed: {ss_response}")
                                else:
                                    st.info("ℹ️ Step 4: No sea service records found in CV")
                            
                            successful_uploads.append({
                                "filename": file.name,
                                "document_id": doc_response.get('id'),
                                "cv_submission_id": cv_response.get('id') if cv_success else None,
                                "user_id": user_id,
                                "fields_extracted": len([v for v in result.values() if v and v != [] and v != ""]),
                                **{k: v for k, v in result.items() if k != 'sea_services'}
                            })
                        else:
                            st.error(f"❌ Document API Error for {file.name}: {doc_response}")
                
                progress_bar.progress((i + 1) / len(uploaded_files))
                
                if i < len(uploaded_files) - 1:
                    status_text.text(f"Pausing 15s for Gemini API limits (File {i+1} of {len(uploaded_files)} done)...")
                    time.sleep(15)
            
            status_text.text("✅ All processing complete!")
            
            if successful_uploads:
                with st.expander("📊 Summary of All Uploads"):
                    st.json(successful_uploads)
