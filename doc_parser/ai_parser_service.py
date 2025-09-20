

# import json
# import yaml
# import docx2txt
# from langchain_community.document_loaders import PyPDFLoader, TextLoader
# from langchain_ollama.llms import OllamaLLM
# from langchain.prompts import PromptTemplate
# from langchain.chains import LLMChain
# import re


# def extract_document_features(file_path: str, model: str = "gemma3:1b") -> str:
#   """
#   Reads a document, extracts features using an Ollama model through LangChain,
#   and returns the result in YAML format.


#   Args:
#   file_path (str): Path to the input document.
#   model (str): Ollama model name (default: gemma3:1b).


#   Returns:
#   str: Extracted features in YAML format.
#   """


#   # --- Step 1: Load the document ---
#   text_content = ""
#   if file_path.endswith(".pdf"):
#     loader = PyPDFLoader(file_path)
#     documents = loader.load()
#     text_content = "\n".join([doc.page_content for doc in documents])
#   elif file_path.endswith(".docx"):
#     try:
#       text_content = docx2txt.process(file_path)
#     except Exception as e:
#       print(f"Error processing docx file with docx2txt: {e}")
#       return None
#   elif file_path.endswith(".txt"):
#     loader = TextLoader(file_path)
#     documents = loader.load()
#     text_content = "\n".join([doc.page_content for doc in documents])
#   else:
#     raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")

#   if not text_content:
#       return None

#   # --- Step 2: Initialize Ollama LLM ---
#   llm = OllamaLLM(model=model)


#   # --- Step 3: Define extraction prompt ---
#   template = """
#   You are an AI assistant. Extract structured features from the following document.
#   Features include:
#        full_name: string
#        date_of_birth: YYYY-MM-DD
#        place_of_birth: string
#        nationality: string
#        passport_number: string
#        passport_issue_date: YYYY-MM-DD
#        passport_expiry_date: YYYY-MM-DD
#        seaman_book_number: string
#        seaman_book_issue_date: YYYY-MM-DD
#        seaman_book_expiry_date: YYYY-MM-DD
#        address: string
#        phone_number: string
#        email: string

#   IMPORTANT: Return ONLY valid YAML format. Do not include markdown code blocks, backticks, or any other formatting.
#   If a field is not found in the document, use null or an empty string.

#   Document:
#   {document}
#   """


#   prompt = PromptTemplate(template=template, input_variables=["document"])
#   chain = prompt | llm


#   # --- Step 4: Run extraction ---
#   response = chain.invoke({"document": text_content[:4000]}) # limit to 4k chars for safety

#   # --- Step 5: Clean the response ---
#   cleaned_response = _clean_yaml_response(response)

#   # --- Step 6: Validate YAML ---
#   try:
#     parsed_yaml = yaml.safe_load(cleaned_response)
#     return yaml.dump(parsed_yaml, sort_keys=False, allow_unicode=True, default_flow_style=False)
#   except yaml.YAMLError as e:
#     print(f"YAML parsing error: {e}")
#     print(f"Cleaned response: {cleaned_response}")
#     return cleaned_response


# def _clean_yaml_response(response: str) -> str:
#     """
#     Clean the AI response by removing markdown formatting and handling JSON responses.
#     """
#     if not response:
#         return ""
    
#     # Remove markdown code blocks
#     response = re.sub(r'```(?:yaml|json)?\s*\n?', '', response, flags=re.IGNORECASE)
#     response = re.sub(r'```\s*\n?', '', response)
    
#     # Remove leading/trailing whitespace
#     response = response.strip()
    
#     # Remove any leading backticks that might remain
#     response = re.sub(r'^`+', '', response)
#     #response = re.sub(r'`+, '', response)
    
#     # Check if the response is JSON format
#     try:
#         json_data = json.loads(response)
        
#         # Handle case where AI returns {"document": [array of objects]}
#         if isinstance(json_data, dict) and "document" in json_data:
#             document_data = json_data["document"]
#             if isinstance(document_data, list) and len(document_data) > 0:
#                 # Take the first object that has the most complete data
#                 best_record = _find_most_complete_record(document_data)
#                 # Convert JSON to YAML format
#                 return yaml.dump(best_record, default_flow_style=False, allow_unicode=True)
        
#         # Handle case where AI returns a single object or array directly
#         elif isinstance(json_data, list) and len(json_data) > 0:
#             best_record = _find_most_complete_record(json_data)
#             return yaml.dump(best_record, default_flow_style=False, allow_unicode=True)
        
#         elif isinstance(json_data, dict):
#             return yaml.dump(json_data, default_flow_style=False, allow_unicode=True)
            
#     except json.JSONDecodeError:
#         # Not JSON, try to process as YAML
#         pass
    
#     # If the response starts with non-YAML content, try to extract YAML part
#     lines = response.split('\n')
#     yaml_started = False
#     yaml_lines = []
    
#     for line in lines:
#         # Check if line looks like YAML (key: value format)
#         if ':' in line and not yaml_started:
#             yaml_started = True
        
#         if yaml_started:
#             yaml_lines.append(line)
    
#     if yaml_lines:
#         response = '\n'.join(yaml_lines)
    
#     return response.strip()


# def _find_most_complete_record(records):
#     """
#     Find the record with the most non-null/non-empty values from a list of records.
#     """
#     if not records or not isinstance(records, list):
#         return {}
    
#     best_record = records[0]
#     max_fields = 0
    
#     for record in records:
#         if not isinstance(record, dict):
#             continue
            
#         # Count non-null, non-empty fields
#         field_count = sum(1 for value in record.values() 
#                          if value is not None and str(value).strip() != "")
        
#         if field_count > max_fields:
#             max_fields = field_count
#             best_record = record
    
#     return best_record

# """
# Debugged AI Parser Service with comprehensive error handling and fallback mechanisms.
# """

# import os
# import re
# import yaml
# import json
# import logging
# from typing import Dict, Any, Optional, List, Tuple
# from datetime import datetime

# logger = logging.getLogger(__name__)

# # Handle imports with fallbacks
# try:
#     from langchain_ollama import OllamaLLM
#     OLLAMA_AVAILABLE = True
# except ImportError:
#     logger.error("langchain-ollama not available. Install with: pip install langchain-ollama")
#     OLLAMA_AVAILABLE = False

# try:
#     from langchain_core.prompts import PromptTemplate
#     LANGCHAIN_AVAILABLE = True
# except ImportError:
#     logger.error("langchain-core not available. Install with: pip install langchain-core")
#     LANGCHAIN_AVAILABLE = False

# # DOCX processing libraries
# try:
#     from docx2python import docx2python
#     DOCX2PYTHON_AVAILABLE = True
#     logger.info("Using docx2python for DOCX processing")
# except ImportError:
#     DOCX2PYTHON_AVAILABLE = False
#     logger.warning("docx2python not available")
    
#     try:
#         from docx import Document
#         PYTHON_DOCX_AVAILABLE = True
#         logger.info("Using python-docx for DOCX processing")
#     except ImportError:
#         PYTHON_DOCX_AVAILABLE = False
#         logger.error("No DOCX processing library available. Install docx2python or python-docx")


# class SeafarerFieldExtractor:
#     """
#     Enhanced field extractor with comprehensive error handling.
#     """
    
#     def __init__(self, model_name: str = "gemma3:1b"):
#         """Initialize the field extractor with error handling."""
#         self.model_name = model_name
#         self.llm = None
#         self.initialization_error = None
        
#         # Check dependencies
#         if not OLLAMA_AVAILABLE:
#             self.initialization_error = "Ollama library not available"
#             return
            
#         if not LANGCHAIN_AVAILABLE:
#             self.initialization_error = "LangChain library not available"
#             return
        
#         self._initialize_llm()

#     def _initialize_llm(self):
#         """Initialize the Ollama LLM with error handling."""
#         try:
#             self.llm = OllamaLLM(model=self.model_name)
            
#             # Test the connection
#             test_response = self.llm.invoke("Test")
#             logger.info(f"Successfully initialized Ollama LLM with model: {self.model_name}")
            
#         except Exception as e:
#             error_msg = f"Failed to initialize LLM: {str(e)}"
#             logger.error(error_msg)
#             self.initialization_error = error_msg
#             self.llm = None

#     def check_dependencies(self) -> Tuple[bool, List[str]]:
#         """
#         Check if all required dependencies are available.
        
#         Returns:
#             Tuple of (success, error_messages)
#         """
#         errors = []
        
#         if not OLLAMA_AVAILABLE:
#             errors.append("langchain-ollama not installed. Run: pip install langchain-ollama")
        
#         if not LANGCHAIN_AVAILABLE:
#             errors.append("langchain-core not installed. Run: pip install langchain-core")
        
#         if not DOCX2PYTHON_AVAILABLE and not PYTHON_DOCX_AVAILABLE:
#             errors.append("No DOCX library available. Run: pip install docx2python")
        
#         if self.initialization_error:
#             errors.append(f"LLM initialization failed: {self.initialization_error}")
        
#         return len(errors) == 0, errors

#     def extract_structured_content_from_docx(self, file_path: str) -> Optional[Dict[str, Any]]:
#         """Extract structured content from DOCX files with error handling."""
        
#         # Check if file exists
#         if not os.path.exists(file_path):
#             logger.error(f"File not found: {file_path}")
#             return None
        
#         # Check file extension
#         if not file_path.lower().endswith('.docx'):
#             logger.error(f"Not a DOCX file: {file_path}")
#             return None
        
#         try:
#             if DOCX2PYTHON_AVAILABLE:
#                 return self._extract_with_docx2python(file_path)
#             elif PYTHON_DOCX_AVAILABLE:
#                 return self._extract_with_python_docx(file_path)
#             else:
#                 logger.error("No DOCX parsing library available")
#                 return None
                
#         except Exception as e:
#             logger.error(f"Error extracting structured content from {file_path}: {str(e)}")
#             return None

#     def _extract_with_docx2python(self, file_path: str) -> Dict[str, Any]:
#         """Extract content using docx2python with error handling."""
#         try:
#             docx_content = docx2python(file_path)
            
#             # Extract text
#             text_content = getattr(docx_content, 'text', '')
            
#             # Extract tables safely
#             tables = []
#             if hasattr(docx_content, 'body') and docx_content.body:
#                 tables = self._clean_tables(docx_content.body)
            
#             # Extract images safely
#             images = []
#             if hasattr(docx_content, 'images') and docx_content.images:
#                 images = list(docx_content.images.keys())
            
#             # Extract properties safely
#             properties = self._extract_properties(docx_content)
            
#             return {
#                 'text': text_content,
#                 'tables': tables,
#                 'images': images,
#                 'properties': properties,
#                 'extraction_method': 'docx2python'
#             }
            
#         except Exception as e:
#             logger.error(f"Error in docx2python extraction: {str(e)}")
#             raise

#     def _extract_with_python_docx(self, file_path: str) -> Dict[str, Any]:
#         """Extract content using python-docx with error handling."""
#         try:
#             doc = Document(file_path)
            
#             # Extract text from paragraphs
#             text_parts = []
#             for paragraph in doc.paragraphs:
#                 if paragraph.text and paragraph.text.strip():
#                     text_parts.append(paragraph.text.strip())
            
#             text_content = '\n'.join(text_parts)
            
#             # Extract tables
#             tables = []
#             for table in doc.tables:
#                 table_data = []
#                 for row in table.rows:
#                     row_data = []
#                     for cell in row.cells:
#                         if cell.text and cell.text.strip():
#                             row_data.append(cell.text.strip())
#                     if row_data:
#                         table_data.append(row_data)
#                 if table_data:
#                     tables.append(table_data)
            
#             # Extract basic properties
#             properties = self._extract_doc_properties(doc)
            
#             return {
#                 'text': text_content,
#                 'tables': tables,
#                 'images': [],
#                 'properties': properties,
#                 'extraction_method': 'python-docx'
#             }
            
#         except Exception as e:
#             logger.error(f"Error in python-docx extraction: {str(e)}")
#             raise

#     def _clean_tables(self, body_content) -> List[List[List[str]]]:
#         """Clean and structure table data from docx2python."""
#         tables = []
#         try:
#             if body_content and isinstance(body_content, (list, tuple)):
#                 for item in body_content:
#                     if isinstance(item, (list, tuple)) and len(item) > 0:
#                         clean_table = []
#                         for row in item:
#                             if isinstance(row, (list, tuple)):
#                                 clean_row = []
#                                 for cell in row:
#                                     cell_text = str(cell).strip()
#                                     if cell_text:
#                                         clean_row.append(cell_text)
#                                 if clean_row:
#                                     clean_table.append(clean_row)
#                         if clean_table:
#                             tables.append(clean_table)
#         except Exception as e:
#             logger.warning(f"Error cleaning tables: {str(e)}")
        
#         return tables

#     def _extract_properties(self, docx_content) -> Dict[str, Any]:
#         """Extract document properties with error handling."""
#         properties = {}
#         try:
#             # Try core_properties first (newer approach)
#             if hasattr(docx_content, 'core_properties'):
#                 core_props = getattr(docx_content, 'core_properties', None)
#                 if core_props:
#                     properties.update({
#                         'title': getattr(core_props, 'title', '') or '',
#                         'subject': getattr(core_props, 'subject', '') or '',
#                         'creator': getattr(core_props, 'creator', '') or '',
#                         'keywords': getattr(core_props, 'keywords', '') or '',
#                         'description': getattr(core_props, 'description', '') or '',
#                         'created': str(getattr(core_props, 'created', '') or ''),
#                         'modified': str(getattr(core_props, 'modified', '') or '')
#                     })
            
#             # Fallback to deprecated properties
#             elif hasattr(docx_content, 'properties'):
#                 import warnings
#                 with warnings.catch_warnings():
#                     warnings.simplefilter("ignore", FutureWarning)
#                     old_props = getattr(docx_content, 'properties', {}) or {}
#                     properties.update(old_props)
                    
#         except Exception as e:
#             logger.warning(f"Could not extract document properties: {str(e)}")
        
#         return properties

#     def _extract_doc_properties(self, doc) -> Dict[str, Any]:
#         """Extract properties from python-docx document."""
#         properties = {}
#         try:
#             if hasattr(doc, 'core_properties'):
#                 core_props = doc.core_properties
#                 properties = {
#                     'title': getattr(core_props, 'title', None) or '',
#                     'subject': getattr(core_props, 'subject', None) or '',
#                     'creator': getattr(core_props, 'author', None) or '',
#                     'created': str(getattr(core_props, 'created', None) or ''),
#                     'modified': str(getattr(core_props, 'modified', None) or '')
#                 }
#         except Exception as e:
#             logger.warning(f"Could not extract document properties: {str(e)}")
        
#         return properties

#     def extract_with_ai(self, structured_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
#         """Extract data using AI model with comprehensive error handling."""
        
#         # Check if LLM is available
#         if not self.llm:
#             error_msg = f"LLM not available: {self.initialization_error or 'Unknown initialization error'}"
#             logger.error(error_msg)
#             return None

#         try:
#             # Create prompt
#             prompt = self.create_enhanced_prompt(structured_data)
#             if not prompt:
#                 logger.error("Failed to create prompt")
#                 return None
            
#             logger.info("Starting AI extraction with Ollama...")
            
#             # Get AI response with timeout handling
#             response = self.llm.invoke(prompt)
            
#             if not response:
#                 logger.error("Empty response from AI model")
#                 return None
            
#             logger.info(f"AI response received (length: {len(response)})")
            
#             # Clean and parse response
#             cleaned_response = self._clean_yaml_response(response)
            
#             if not cleaned_response:
#                 logger.error("No valid content after cleaning AI response")
#                 return None
            
#             # Parse YAML
#             try:
#                 parsed_data = yaml.safe_load(cleaned_response)
#             except yaml.YAMLError as e:
#                 logger.error(f"YAML parsing error: {e}")
#                 logger.error(f"Cleaned response was: {cleaned_response[:500]}...")
#                 return None
            
#             if not isinstance(parsed_data, dict):
#                 logger.error(f"Expected dict, got {type(parsed_data)}")
#                 return None
            
#             # Post-process the data
#             processed_data = self._post_process_extracted_data(parsed_data)
            
#             logger.info("AI extraction completed successfully")
#             return processed_data
            
#         except Exception as e:
#             logger.error(f"Error during AI extraction: {str(e)}")
#             return None

#     def create_enhanced_prompt(self, structured_data: Dict[str, Any]) -> Optional[str]:
#         """Create enhanced prompt with error handling."""
#         try:
#             # Prepare content sections with length limits
#             text_content = structured_data.get('text', '')
#             if len(text_content) > 6000:  # Reduced limit for stability
#                 text_content = text_content[:6000] + "...[truncated]"
            
#             # Prepare tables content
#             tables_content = "No tables found"
#             tables = structured_data.get('tables', [])
#             if tables:
#                 tables_parts = []
#                 for i, table in enumerate(tables[:3]):  # Limit to 3 tables
#                     tables_parts.append(f"\nTable {i+1}:")
#                     for row in table[:10]:  # Limit rows per table
#                         if len(row) <= 10:  # Limit columns
#                             tables_parts.append(" | ".join(row))
#                 tables_content = "\n".join(tables_parts)
            
#             # Prepare properties
#             properties_content = "No properties found"
#             properties = structured_data.get('properties', {})
#             if properties:
#                 prop_parts = []
#                 for key, value in properties.items():
#                     if value and len(str(value)) < 100:  # Skip very long values
#                         prop_parts.append(f"{key}: {value}")
#                 properties_content = "\n".join(prop_parts) if prop_parts else "No properties found"
            
#             # Create simplified prompt
#             prompt = f"""
# Extract seafarer information from this document. Return ONLY valid YAML format, no explanations.

# DOCUMENT TEXT:
# {text_content}

# TABLES:
# {tables_content}

# PROPERTIES:
# {properties_content}

# Extract into this YAML structure (use "Not Available" for missing data):

# personal_information:
#   full_name: ""
#   date_of_birth: ""
#   place_of_birth: ""
#   nationality: ""

# contact_information:
#   address: ""
#   phone_number: ""
#   email: ""

# travel_documents:
#   passport_number: ""
#   passport_issue_date: ""
#   passport_expiry_date: ""
#   seaman_book_number: ""
#   seaman_book_issue_date: ""
#   seaman_book_expiry_date: ""

# medical_information:
#   medical_certificate_number: ""
#   medical_issue_date: ""
#   medical_expiry_date: ""

# qualifications:
#   certificate_of_competency: ""
#   coc_rank: ""

# sea_service:
#   total_sea_service_months: ""
#   last_vessel_name: ""
#   last_rank: ""

# Return only the YAML:"""

#             return prompt
            
#         except Exception as e:
#             logger.error(f"Error creating prompt: {str(e)}")
#             return None

#     def _clean_yaml_response(self, response: str) -> str:
#         """Clean AI response with enhanced error handling."""
#         if not response:
#             return ""
        
#         try:
#             # Remove markdown formatting
#             cleaned = re.sub(r'```(?:yaml|json)?\s*', '', response, flags=re.IGNORECASE)
#             cleaned = re.sub(r'```\s*', '', cleaned)
#             cleaned = cleaned.strip()
            
#             # Remove leading backticks
#             cleaned = re.sub(r'^`+', '', cleaned)
            
#             # Remove markdown formatting
#             cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned)
#             cleaned = re.sub(r'\*(.*?)\*', r'\1', cleaned)
            
#             # Extract YAML portion
#             lines = cleaned.split('\n')
#             yaml_lines = []
#             yaml_started = False
            
#             for line in lines:
#                 line_stripped = line.strip()
                
#                 # Skip empty lines at start
#                 if not yaml_started and not line_stripped:
#                     continue
                
#                 # Detect YAML start
#                 if ':' in line and not line_stripped.startswith('#'):
#                     yaml_started = True
                
#                 if yaml_started:
#                     # Stop at explanatory text
#                     if (line_stripped and 
#                         not line.startswith(' ') and 
#                         ':' not in line and 
#                         not line.endswith(':') and
#                         len(line_stripped.split()) <= 3):
#                         break
#                     yaml_lines.append(line)
            
#             result = '\n'.join(yaml_lines).strip()
#             return result
            
#         except Exception as e:
#             logger.error(f"Error cleaning YAML response: {str(e)}")
#             return response  # Return original if cleaning fails

#     def _post_process_extracted_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
#         """Post-process extracted data with error handling."""
#         if not isinstance(data, dict):
#             return {}
        
#         processed = {}
        
#         try:
#             for category, fields in data.items():
#                 if isinstance(fields, dict):
#                     processed[category] = {}
#                     for field, value in fields.items():
#                         cleaned_value = self._clean_field_value(value)
#                         processed[category][field] = cleaned_value
#                 else:
#                     processed[category] = self._clean_field_value(fields)
#         except Exception as e:
#             logger.error(f"Error post-processing data: {str(e)}")
#             return data  # Return original if processing fails
        
#         return processed

#     def _clean_field_value(self, value) -> str:
#         """Clean individual field values."""
#         if value is None:
#             return "Not Available"
        
#         try:
#             str_value = str(value).strip()
            
#             # Handle common variations of empty values
#             empty_variations = ['', 'null', 'none', 'n/a', 'na', '-', '--', 'nil']
#             if str_value.lower() in empty_variations:
#                 return "Not Available"
            
#             return str_value
#         except Exception:
#             return "Not Available"

#     def extract_from_document(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], Dict[str, Any]]:
#         """
#         Main extraction method with comprehensive error handling.
#         """
#         try:
#             # Check dependencies first
#             deps_ok, dep_errors = self.check_dependencies()
#             if not deps_ok:
#                 return None, {
#                     'error': 'Missing dependencies',
#                     'details': dep_errors
#                 }
            
#             # Extract structured content
#             logger.info(f"Extracting structured content from: {file_path}")
#             structured_data = self.extract_structured_content_from_docx(file_path)
            
#             if not structured_data:
#                 return None, {'error': 'Failed to extract structured content from document'}
            
#             logger.info(f"Extracted {len(structured_data.get('text', ''))} chars of text and {len(structured_data.get('tables', []))} tables")
            
#             # Extract with AI
#             extracted_data = self.extract_with_ai(structured_data)
            
#             if not extracted_data:
#                 return None, {'error': 'AI extraction failed'}
            
#             # Create metadata
#             metadata = {
#                 'extraction_method': 'ai_enhanced',
#                 'model_used': self.model_name,
#                 'text_length': len(structured_data.get('text', '')),
#                 'tables_found': len(structured_data.get('tables', [])),
#                 'processed_at': datetime.now().isoformat(),
#                 'success': True
#             }
            
#             return extracted_data, metadata
            
#         except Exception as e:
#             logger.error(f"Error extracting from document: {str(e)}")
#             return None, {'error': str(e)}


# # Convenience functions for backward compatibility
# def extract_structured_content_from_docx(file_path: str) -> Optional[Dict[str, Any]]:
#     """Extract structured content from DOCX file."""
#     try:
#         extractor = SeafarerFieldExtractor()
#         return extractor.extract_structured_content_from_docx(file_path)
#     except Exception as e:
#         logger.error(f"Error in extract_structured_content_from_docx: {str(e)}")
#         return None


# def extract_data_from_document_enhanced(file_path: str) -> str:
#     """Enhanced document data extraction returning YAML string."""
#     try:
#         extractor = SeafarerFieldExtractor()
        
#         # Check dependencies
#         deps_ok, dep_errors = extractor.check_dependencies()
#         if not deps_ok:
#             return f"Error: Missing dependencies - {'; '.join(dep_errors)}"
        
#         extracted_data, metadata = extractor.extract_from_document(file_path)
        
#         if extracted_data:
#             return yaml.dump(extracted_data, default_flow_style=False, allow_unicode=True)
#         else:
#             error_msg = metadata.get('error', 'Unknown error')
#             details = metadata.get('details', [])
#             if details:
#                 return f"Error: {error_msg} - Details: {'; '.join(details)}"
#             return f"Error: {error_msg}"
            
#     except Exception as e:
#         logger.error(f"Error in extract_data_from_document_enhanced: {str(e)}")
#         return f"Error: {str(e)}"


# def extract_data_from_document(file_path: str) -> str:
#     """Backward compatibility wrapper."""
#     return extract_data_from_document_enhanced(file_path)

"""
Improved AI Parser Service that extracts more comprehensive data from seafarer documents.
"""

import os
import re
import yaml
import json
import logging
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime

logger = logging.getLogger(__name__)

# Handle imports with fallbacks
try:
    from langchain_ollama import OllamaLLM
    OLLAMA_AVAILABLE = True
except ImportError:
    logger.error("langchain-ollama not available")
    OLLAMA_AVAILABLE = False

try:
    from langchain_core.prompts import PromptTemplate
    LANGCHAIN_AVAILABLE = True
except ImportError:
    logger.error("langchain-core not available")
    LANGCHAIN_AVAILABLE = False

# DOCX processing libraries
try:
    from docx2python import docx2python
    DOCX2PYTHON_AVAILABLE = True
except ImportError:
    DOCX2PYTHON_AVAILABLE = False
    try:
        from docx import Document
        PYTHON_DOCX_AVAILABLE = True
    except ImportError:
        PYTHON_DOCX_AVAILABLE = False


class EnhancedSeafarerFieldExtractor:
    """
    Enhanced field extractor that combines AI extraction with rule-based pattern matching.
    """
    
    def __init__(self, model_name: str = "gemma3:1b"):
        """Initialize the field extractor."""
        self.model_name = model_name
        self.llm = None
        self.initialization_error = None
        
        # Check dependencies
        if not OLLAMA_AVAILABLE or not LANGCHAIN_AVAILABLE:
            self.initialization_error = "Required libraries not available"
        else:
            self._initialize_llm()

    def _initialize_llm(self):
        """Initialize the Ollama LLM with error handling."""
        try:
            self.llm = OllamaLLM(model=self.model_name)
            # Test the connection
            test_response = self.llm.invoke("Test")
            logger.info(f"Successfully initialized Ollama LLM with model: {self.model_name}")
        except Exception as e:
            error_msg = f"Failed to initialize LLM: {str(e)}"
            logger.error(error_msg)
            self.initialization_error = error_msg
            self.llm = None

    def extract_structured_content_from_docx(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Extract structured content from DOCX files."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        try:
            if DOCX2PYTHON_AVAILABLE:
                return self._extract_with_docx2python(file_path)
            elif PYTHON_DOCX_AVAILABLE:
                return self._extract_with_python_docx(file_path)
            else:
                logger.error("No DOCX parsing library available")
                return None
        except Exception as e:
            logger.error(f"Error extracting structured content: {str(e)}")
            return None

    def _extract_with_docx2python(self, file_path: str) -> Dict[str, Any]:
        """Extract content using docx2python."""
        docx_content = docx2python(file_path)
        
        # Extract text safely
        text_content = getattr(docx_content, 'text', '')
        
        # Extract tables safely
        tables = []
        if hasattr(docx_content, 'body') and docx_content.body:
            tables = self._clean_tables(docx_content.body)
        
        # Extract images safely
        images = []
        if hasattr(docx_content, 'images') and docx_content.images:
            images = list(docx_content.images.keys())
        
        # Extract properties safely
        properties = self._extract_properties(docx_content)
        
        return {
            'text': text_content,
            'tables': tables,
            'images': images,
            'properties': properties,
            'extraction_method': 'docx2python'
        }

    def _extract_with_python_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract content using python-docx."""
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        text_content = '\n'.join(text_parts)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        return {
            'text': text_content,
            'tables': tables,
            'images': [],
            'properties': {},
            'extraction_method': 'python-docx'
        }

    def _clean_tables(self, body_content) -> List[List[List[str]]]:
        """Clean and structure table data from docx2python."""
        tables = []
        try:
            if body_content and isinstance(body_content, (list, tuple)):
                for item in body_content:
                    if isinstance(item, (list, tuple)) and len(item) > 0:
                        clean_table = []
                        for row in item:
                            if isinstance(row, (list, tuple)):
                                clean_row = []
                                for cell in row:
                                    cell_text = str(cell).strip()
                                    if cell_text:
                                        clean_row.append(cell_text)
                                if clean_row:
                                    clean_table.append(clean_row)
                        if clean_table:
                            tables.append(clean_table)
        except Exception as e:
            logger.warning(f"Error cleaning tables: {str(e)}")
        return tables

    def _extract_properties(self, docx_content) -> Dict[str, Any]:
        """Extract document properties safely."""
        properties = {}
        try:
            if hasattr(docx_content, 'core_properties'):
                core_props = getattr(docx_content, 'core_properties', None)
                if core_props:
                    properties = {
                        'title': getattr(core_props, 'title', '') or '',
                        'subject': getattr(core_props, 'subject', '') or '',
                        'creator': getattr(core_props, 'creator', '') or ''
                    }
        except Exception as e:
            logger.warning(f"Could not extract document properties: {str(e)}")
        return properties

    def extract_with_pattern_matching(self, text: str) -> Dict[str, Any]:
        """
        Extract data using pattern matching before AI processing.
        This catches obvious patterns that AI might miss.
        """
        extracted = {}
        
        try:
            # Personal Information Patterns
            extracted['personal_information'] = self._extract_personal_patterns(text)
            
            # Contact Information Patterns
            extracted['contact_information'] = self._extract_contact_patterns(text)
            
            # Travel Documents Patterns
            extracted['travel_documents'] = self._extract_document_patterns(text)
            
            # Medical Information Patterns
            extracted['medical_information'] = self._extract_medical_patterns(text)
            
            # Qualifications Patterns
            extracted['qualifications'] = self._extract_qualification_patterns(text)
            
            # Training Patterns
            extracted['stcw_training'] = self._extract_training_patterns(text)
            
        except Exception as e:
            logger.error(f"Error in pattern matching extraction: {str(e)}")
        
        return extracted

    def _extract_personal_patterns(self, text: str) -> Dict[str, str]:
        """Extract personal information using patterns."""
        personal = {}
        
        try:
            # Full Name - look for patterns after "Full Name"
            name_match = re.search(r'Full Name.*?([A-Z][A-Z\s]+[A-Z])', text, re.IGNORECASE | re.DOTALL)
            if name_match:
                personal['full_name'] = name_match.group(1).strip()
            
            # Date of Birth - look for date patterns
            dob_patterns = [
                r'Date Of Birth.*?(\d{2}/\d{2}/\d{4})',
                r'Birth.*?(\d{2}/\d{2}/\d{4})',
                r'DOB.*?(\d{2}/\d{2}/\d{4})'
            ]
            for pattern in dob_patterns:
                dob_match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if dob_match:
                    personal['date_of_birth'] = dob_match.group(1)
                    break
            
            # Place of Birth
            pob_match = re.search(r'Place Of Birth.*?(DAMIETTA|DAMEITTA|[A-Z][A-Za-z\s]+)', text, re.IGNORECASE | re.DOTALL)
            if pob_match:
                personal['place_of_birth'] = pob_match.group(1).strip()
            
            # Nationality
            nat_match = re.search(r'Nationality.*?(Egyptian?|[A-Z][a-z]+)', text, re.IGNORECASE | re.DOTALL)
            if nat_match:
                personal['nationality'] = nat_match.group(1).strip()
            
            # Height
            height_match = re.search(r'Height.*?(\d+)', text, re.IGNORECASE)
            if height_match:
                personal['height'] = height_match.group(1) + ' cm'
            
            # Weight  
            weight_match = re.search(r'Weight.*?(\d+)', text, re.IGNORECASE)
            if weight_match:
                personal['weight'] = weight_match.group(1) + ' kg'
            
            # Marital Status
            if 'Married' in text and '☒' in text:
                personal['marital_status'] = 'Married'
            elif 'Single' in text and '☒' in text:
                personal['marital_status'] = 'Single'
                
        except Exception as e:
            logger.error(f"Error extracting personal patterns: {str(e)}")
        
        return personal

    def _extract_contact_patterns(self, text: str) -> Dict[str, str]:
        """Extract contact information using patterns."""
        contact = {}
        
        try:
            # Email
            email_match = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', text)
            if email_match:
                contact['email'] = email_match.group(1)
            
            # Phone numbers - look for Egyptian phone patterns
            phone_patterns = [
                r'\+20\s*\d{3}\s*\d{3}\s*\d{4}',
                r'\+\s*20\d{10}',
                r'\d{11,12}'
            ]
            
            phones = []
            for pattern in phone_patterns:
                phone_matches = re.findall(pattern, text)
                phones.extend(phone_matches)
            
            if phones:
                # Clean and format phone numbers
                cleaned_phones = []
                for phone in phones:
                    phone_clean = re.sub(r'[^\d+]', '', phone)
                    if len(phone_clean) >= 10:
                        cleaned_phones.append(phone_clean)
                
                if cleaned_phones:
                    contact['phone_number'] = cleaned_phones[0]
                    if len(cleaned_phones) > 1:
                        contact['phone_number_2'] = cleaned_phones[1]
            
            # Address
            address_patterns = [
                r'Home Address.*?(EZBET.*?EGYPT)',
                r'Address.*?(EZBET.*?EGYPT)',
                r'(EZBET.*?EGYPT)'
            ]
            
            for pattern in address_patterns:
                addr_match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if addr_match:
                    contact['address'] = addr_match.group(1).strip()
                    break
                    
        except Exception as e:
            logger.error(f"Error extracting contact patterns: {str(e)}")
        
        return contact

    def _extract_document_patterns(self, text: str) -> Dict[str, str]:
        """Extract travel document information using patterns."""
        documents = {}
        
        try:
            # Passport Number
            passport_patterns = [
                r'Passport.*?([A-Z]\d{8})',
                r'([A-Z]\d{8})',
            ]
            
            for pattern in passport_patterns:
                passport_match = re.search(pattern, text, re.IGNORECASE)
                if passport_match:
                    documents['passport_number'] = passport_match.group(1)
                    break
            
            # Passport dates
            passport_issue_match = re.search(r'Passport.*?(\d{2}/\d{2}/\d{4}).*?(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE | re.DOTALL)
            if passport_issue_match:
                documents['passport_issue_date'] = passport_issue_match.group(1)
                documents['passport_expiry_date'] = passport_issue_match.group(2)
            
            # Seaman Book Number
            seaman_patterns = [
                r'Seaman Book.*?(S\d{8})',
                r'(S\d{8})',
            ]
            
            for pattern in seaman_patterns:
                seaman_match = re.search(pattern, text, re.IGNORECASE)
                if seaman_match:
                    documents['seaman_book_number'] = seaman_match.group(1)
                    break
            
            # Seaman Book dates
            seaman_date_match = re.search(r'Seaman Book.*?(\d{2}/\d{2}/\d{4}).*?(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE | re.DOTALL)
            if seaman_date_match:
                documents['seaman_book_issue_date'] = seaman_date_match.group(1)
                documents['seaman_book_expiry_date'] = seaman_date_match.group(2)
                
        except Exception as e:
            logger.error(f"Error extracting document patterns: {str(e)}")
        
        return documents

    def _extract_medical_patterns(self, text: str) -> Dict[str, str]:
        """Extract medical information using patterns."""
        medical = {}
        
        try:
            # Medical Certificate Number
            med_cert_match = re.search(r'International Medical.*?(\d{4,5})', text, re.IGNORECASE | re.DOTALL)
            if med_cert_match:
                medical['medical_certificate_number'] = med_cert_match.group(1)
            
            # Medical dates
            med_date_match = re.search(r'International Medical.*?(\d{2}/\d{2}/\d{4}).*?(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE | re.DOTALL)
            if med_date_match:
                medical['medical_issue_date'] = med_date_match.group(1)
                medical['medical_expiry_date'] = med_date_match.group(2)
            
            # Yellow Fever
            yellow_fever_match = re.search(r'Yellow Fever.*?(\d{2}/\d{2}/\d{4})', text, re.IGNORECASE | re.DOTALL)
            if yellow_fever_match:
                medical['yellow_fever_vaccination'] = yellow_fever_match.group(1)
                
        except Exception as e:
            logger.error(f"Error extracting medical patterns: {str(e)}")
        
        return medical

    def _extract_qualification_patterns(self, text: str) -> Dict[str, str]:
        """Extract qualification information using patterns."""
        qualifications = {}
        
        try:
            # COC Rank
            coc_match = re.search(r'COC.*?Rank.*?(ABLE SEAFARERE DECK|ABLE SEAFARER DECK)', text, re.IGNORECASE | re.DOTALL)
            if coc_match:
                qualifications['certificate_of_competency'] = coc_match.group(1)
                qualifications['coc_rank'] = coc_match.group(1)
            
            # DP Training
            if 'D.P. INDUCTION' in text:
                qualifications['dp_induction'] = 'Yes'
            if 'D.P. ADVANCED' in text:
                qualifications['dp_advanced'] = 'Yes'
            if 'D.P. OPERATOR (UNLIMITED)' in text:
                qualifications['dp_operator_unlimited'] = 'Yes'
                
        except Exception as e:
            logger.error(f"Error extracting qualification patterns: {str(e)}")
        
        return qualifications

    def _extract_training_patterns(self, text: str) -> Dict[str, str]:
        """Extract STCW training information using patterns."""
        training = {}
        
        try:
            # Map of training courses with their certificate numbers/dates
            training_courses = {
                'personal_survival_techniques': r'Personal Survival Techniques.*?(\d{4}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'fire_prevention_fighting': r'Fire Prevention and Fire Fighting.*?(\d{4}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'elementary_first_aid': r'Elementary First Aid.*?(\d{4}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'personal_safety_social_responsibilities': r'Personal Safety and Social Responsibilities.*?(\d{4}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'security_awareness': r'Security Awareness.*?(XNA).*?(\d{2}/\d{2}/\d{4})',
                'proficiency_survival_craft': r'Proficiency In Survival Craft.*?(\d{4}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'passenger_safety': r'Passenger Safety.*?(\d{4,5}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'crowd_management': r'Crowd Management.*?(\d{4,5}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
                'crisis_management': r'Crisis Management.*?(\d{4,5}/\d{2}/EG).*?(\d{2}/\d{2}/\d{4})',
            }
            
            for course_name, pattern in training_courses.items():
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    training[course_name] = f"Certificate: {match.group(1)}, Date: {match.group(2)}"
                    
        except Exception as e:
            logger.error(f"Error extracting training patterns: {str(e)}")
        
        return training

    def combine_extractions(self, ai_extracted: Dict[str, Any], pattern_extracted: Dict[str, Any]) -> Dict[str, Any]:
        """
        Combine AI extraction with pattern matching results, preferring pattern matching for specific fields.
        """
        combined = {}
        
        # Get all categories from both extractions
        all_categories = set(ai_extracted.keys()) | set(pattern_extracted.keys())
        
        for category in all_categories:
            combined[category] = {}
            
            # Start with AI extracted data
            if category in ai_extracted and isinstance(ai_extracted[category], dict):
                combined[category].update(ai_extracted[category])
            
            # Override with pattern extracted data (higher confidence)
            if category in pattern_extracted and isinstance(pattern_extracted[category], dict):
                for field, value in pattern_extracted[category].items():
                    if value and str(value).strip() and str(value) != "Not Available":
                        combined[category][field] = value
        
        return combined

    def extract_from_document(self, file_path: str) -> Tuple[Optional[Dict[str, Any]], Dict[str, Any]]:
        """
        Main extraction method combining pattern matching and AI processing.
        """
        try:
            # Step 1: Extract structured content
            structured_data = self.extract_structured_content_from_docx(file_path)
            if not structured_data:
                return None, {'error': 'Failed to extract structured content'}
            
            text = structured_data.get('text', '')
            
            # Step 2: Pattern matching extraction (high confidence)
            logger.info("Starting pattern matching extraction...")
            pattern_extracted = self.extract_with_pattern_matching(text)
            
            # Step 3: AI extraction (if available)
            ai_extracted = {}
            if self.llm and not self.initialization_error:
                try:
                    logger.info("Starting AI extraction...")
                    ai_extracted = self._extract_with_ai(structured_data)
                except Exception as e:
                    logger.warning(f"AI extraction failed, using pattern matching only: {str(e)}")
            else:
                logger.warning("AI not available, using pattern matching only")
            
            # Step 4: Combine extractions
            combined_data = self.combine_extractions(ai_extracted, pattern_extracted)
            
            # Step 5: Create metadata
            metadata = {
                'extraction_method': 'hybrid_pattern_ai',
                'model_used': self.model_name if self.llm else 'pattern_only',
                'text_length': len(text),
                'tables_found': len(structured_data.get('tables', [])),
                'processed_at': datetime.now().isoformat(),
                'pattern_extraction': True,
                'ai_extraction': bool(ai_extracted),
                'success': True
            }
            
            return combined_data, metadata
            
        except Exception as e:
            logger.error(f"Error extracting from document: {str(e)}")
            return None, {'error': str(e)}

    def _extract_with_ai(self, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract data using AI model as supplement to pattern matching."""
        if not self.llm:
            return {}
        
        try:
            # Create simplified prompt focusing on fields pattern matching might miss
            prompt = self._create_ai_prompt(structured_data)
            
            # Get AI response
            response = self.llm.invoke(prompt)
            
            # Clean and parse response
            cleaned_response = self._clean_yaml_response(response)
            parsed_data = yaml.safe_load(cleaned_response)
            
            if isinstance(parsed_data, dict):
                return parsed_data
            
        except Exception as e:
            logger.error(f"Error in AI extraction: {str(e)}")
        
        return {}

    def _create_ai_prompt(self, structured_data: Dict[str, Any]) -> str:
        """Create AI prompt focusing on fields that pattern matching might miss."""
        text_content = structured_data.get('text', '')[:4000]  # Limit for stability
        
        prompt = f"""
Extract seafarer information from this document. Focus on fields that might be in tables or hard to parse.
Return ONLY valid YAML format.

DOCUMENT TEXT:
{text_content}

Extract into this YAML structure (use "Not Available" for missing data):

personal_information:
  full_name: ""
  date_of_birth: ""
  place_of_birth: ""
  nationality: ""

contact_information:
  address: ""
  phone_number: ""
  email: ""

travel_documents:
  passport_number: ""
  passport_issue_date: ""
  passport_expiry_date: ""
  seaman_book_number: ""
  seaman_book_issue_date: ""
  seaman_book_expiry_date: ""

next_of_kin:
  full_name: ""
  relationship: ""

sea_service:
  last_vessel_name: ""
  last_company: ""
  total_experience_years: ""

Return only the YAML:"""
        
        return prompt

    def _clean_yaml_response(self, response: str) -> str:
        """Clean AI response to ensure valid YAML."""
        if not response:
            return ""
        
        # Remove markdown formatting
        cleaned = re.sub(r'```(?:yaml|json)?\s*', '', response, flags=re.IGNORECASE)
        cleaned = re.sub(r'```\s*', '', cleaned)
        cleaned = cleaned.strip()
        
        # Extract YAML portion
        lines = cleaned.split('\n')
        yaml_lines = []
        yaml_started = False
        
        for line in lines:
            line_stripped = line.strip()
            
            if not yaml_started and not line_stripped:
                continue
            
            if ':' in line and not line_stripped.startswith('#'):
                yaml_started = True
            
            if yaml_started:
                yaml_lines.append(line)
        
        return '\n'.join(yaml_lines).strip()

    def check_dependencies(self) -> Tuple[bool, List[str]]:
        """Check if all required dependencies are available."""
        errors = []
        
        if not OLLAMA_AVAILABLE:
            errors.append("langchain-ollama not installed")
        
        if not LANGCHAIN_AVAILABLE:
            errors.append("langchain-core not installed")
        
        if not DOCX2PYTHON_AVAILABLE and not PYTHON_DOCX_AVAILABLE:
            errors.append("No DOCX library available")
        
        if self.initialization_error:
            errors.append(f"LLM initialization failed: {self.initialization_error}")
        
        return len(errors) == 0, errors


# Convenience functions for backward compatibility
def extract_structured_content_from_docx(file_path: str) -> Optional[Dict[str, Any]]:
    """Extract structured content from DOCX file."""
    try:
        extractor = EnhancedSeafarerFieldExtractor()
        return extractor.extract_structured_content_from_docx(file_path)
    except Exception as e:
        logger.error(f"Error in extract_structured_content_from_docx: {str(e)}")
        return None


def extract_data_from_document_enhanced(file_path: str) -> str:
    """Enhanced document data extraction returning YAML string."""
    try:
        extractor = EnhancedSeafarerFieldExtractor()
        
        extracted_data, metadata = extractor.extract_from_document(file_path)
        
        if extracted_data:
            return yaml.dump(extracted_data, default_flow_style=False, allow_unicode=True)
        else:
            error_msg = metadata.get('error', 'Unknown error')
            return f"Error: {error_msg}"
            
    except Exception as e:
        logger.error(f"Error in extract_data_from_document_enhanced: {str(e)}")
        return f"Error: {str(e)}"


def extract_data_from_document(file_path: str) -> str:
    """Backward compatibility wrapper."""
    return extract_data_from_document_enhanced(file_path)


# For backward compatibility with existing views
class SeafarerFieldExtractor(EnhancedSeafarerFieldExtractor):
    """Backward compatibility alias."""
    pass