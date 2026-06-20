import docx
import json
import sys
import re

def clean(text):
    if text is None: return ""
    return text.strip().replace('\n', ' ')

def extract_from_docx(file_path):
    doc = docx.Document(file_path)
    data = {
        "user_name": "",
        "user_email_display": "",
        "position_name": "",
        "status": "Pending",
        "notes": "Auto-submitted via local DOCX parser",
        "seafarer_application": {
            "1_personal_details": {},
            "3_contact_details": {},
            "4_travel_documents": [],
            "5_professional_qualification_certificate_of_competency": [],
            "8_marine_courses": [],
            "9_complete_sea_service_details": {"applicant_info": {}, "service_records": []}
        },
        "user_documents": {
            "passport": {},
            "seaman_book": {},
            "coc": {},
            "goc": {},
            "health_certificate": {},
            "sea_service": [],
            "marine_courses": []
        }
    }

    # Helper to deduplicate row cells (due to merged cells)
    def get_row_cells(row):
        cells = []
        for cell in row.cells:
            t = clean(cell.text)
            if not cells or cells[-1] != t:
                cells.append(t)
        return cells

    for table in doc.tables:
        text_content = " ".join([clean(c.text) for r in table.rows for c in r.cells])
        
        # Application info
        if "Application For Position" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                if len(cells) >= 2:
                    if "Application For Position" in cells[0]:
                        data["position_name"] = cells[1]

        # Personal details
        if "PERSONAL DETAILS" in text_content and "Full Name" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                for i, c in enumerate(cells):
                    if c.startswith("Full Name") and i+1 < len(cells):
                        data["user_name"] = cells[i+1]
                    if c.startswith("Date Of Birth") and i+1 < len(cells):
                        data["seafarer_application"]["1_personal_details"]["date_of_birth"] = cells[i+1]
        
        # Contact Details
        if "3. CONTACT DETAILS" in text_content:
            for row in table.rows:
                cells = get_row_cells(row)
                for i, c in enumerate(cells):
                    if "E-mail" in c and i+1 < len(cells):
                        data["user_email_display"] = cells[i+1]
                    if "Mobile / Tel" in c and i+1 < len(cells):
                        data["seafarer_application"]["3_contact_details"]["mobile_tel"] = cells[i+1]

        # Travel Documents
        if "TRAVEL DOCUMENTS" in text_content and "Passport" in text_content:
            for row in table.rows[2:]: # Skip headers
                cells = get_row_cells(row)
                if len(cells) >= 5:
                    doc_type = cells[0]
                    if "Passport" in doc_type:
                        data["user_documents"]["passport"] = {
                            "passport_no": cells[1], "issue_date": cells[2], "expiry_date": cells[3]
                        }
                    elif "Seaman" in doc_type:
                        data["user_documents"]["seaman_book"] = {
                            "seaman_book_no": cells[1], "issue_date": cells[2], "expiry_date": cells[3]
                        }

        # Courses
        if "MARINE COURSES" in text_content:
            for row in table.rows[2:]:
                cells = get_row_cells(row)
                if len(cells) >= 3 and cells[0] and cells[0] != "Course Name":
                    data["user_documents"]["marine_courses"].append({
                        "course_name": cells[0],
                        "number": cells[1] if len(cells)>1 else "",
                        "issue_date": cells[2] if len(cells)>2 else "",
                        "expiry_date": cells[3] if len(cells)>3 else ""
                    })

        # Sea Service
        if "COMPLETE SEA" in text_content:
            for row in table.rows[2:]:
                cells = get_row_cells(row)
                if len(cells) >= 5 and cells[0] and cells[0] != "Company Name":
                    data["user_documents"]["sea_service"].append({
                        "company_name": cells[0],
                        "rank": cells[1] if len(cells)>1 else "",
                        "vessel_name": cells[2] if len(cells)>2 else "",
                        "signed_on": cells[4] if len(cells)>4 else "",
                        "signed_off": cells[5] if len(cells)>5 else ""
                    })

    # Find email in paragraphs if not found
    if not data["user_email_display"]:
        for p in doc.paragraphs:
            match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', p.text)
            if match:
                data["user_email_display"] = match.group(0)

    print(json.dumps(data, indent=2))

if __name__ == "__main__":
    extract_from_docx(sys.argv[1])
