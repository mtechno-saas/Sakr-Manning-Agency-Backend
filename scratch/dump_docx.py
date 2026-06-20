import docx
import sys

def dump(file_path):
    doc = docx.Document(file_path)
    print("=== PARAGRAPHS ===")
    for p in doc.paragraphs[:10]:
        print(repr(p.text))
    print("\n=== TABLES ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows[:3]:
            print(" | ".join(repr(cell.text) for cell in row.cells))

if __name__ == "__main__":
    dump(sys.argv[1])
